from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'downloads'
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / 'Casey_Barton_Resume.docx'

NAVY='0B2028'; TEAL='087A73'; MINT='7DEAD4'; LIGHT='EEF6F4'; LIGHT2='F7FAF9'; MID='D4E4E1'; DARK='14232A'; MUTED='5B6D75'; AMBER='A56600'; AMBER_BG='FFF5DE'; WHITE='FFFFFF'
FONT='Carlito'


def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None:
        shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def borders(cell, color=MID, size='2'):
    tcPr=cell._tc.get_or_add_tcPr(); node=tcPr.find(qn('w:tcBorders'))
    if node is None: node=OxmlElement('w:tcBorders'); tcPr.append(node)
    for edge in ('top','left','bottom','right'):
        e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),size); e.set(qn('w:color'),color); node.append(e)

def cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.find(qn('w:tcMar'))
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for n,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        x=OxmlElement(f'w:{n}'); x.set(qn('w:w'),str(v)); x.set(qn('w:type'),'dxa'); tcMar.append(x)

def no_spacing(table):
    tblPr=table._tbl.tblPr; x=tblPr.find(qn('w:tblCellSpacing'))
    if x is None: x=OxmlElement('w:tblCellSpacing'); tblPr.append(x)
    x.set(qn('w:w'),'0'); x.set(qn('w:type'),'dxa')

def run(p,text,bold=False,color=DARK,size=9,italic=False):
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.name=FONT; r._element.rPr.rFonts.set(qn('w:eastAsia'),FONT); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); return r

def para(p,before=0,after=0,line=1.0):
    f=p.paragraph_format; f.space_before=Pt(before); f.space_after=Pt(after); f.line_spacing=line

def add_link(p,text,url,color=TEAL,size=8.5):
    part=p.part; rid=part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True)
    h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); rr=OxmlElement('w:r'); rpr=OxmlElement('w:rPr')
    c=OxmlElement('w:color'); c.set(qn('w:val'),color); rpr.append(c); u=OxmlElement('w:u'); u.set(qn('w:val'),'none'); rpr.append(u)
    sz=OxmlElement('w:sz'); sz.set(qn('w:val'),str(int(size*2))); rpr.append(sz); szc=OxmlElement('w:szCs'); szc.set(qn('w:val'),str(int(size*2))); rpr.append(szc)
    fonts=OxmlElement('w:rFonts'); fonts.set(qn('w:ascii'),FONT); fonts.set(qn('w:hAnsi'),FONT); rpr.append(fonts)
    rr.append(rpr); t=OxmlElement('w:t'); t.text=text; rr.append(t); h.append(rr); p._p.append(h)

def section_title(doc,kicker,title):
    p=doc.add_paragraph(); para(p,before=7,after=1); run(p,kicker.upper(),True,TEAL,7.3)
    p=doc.add_paragraph(); para(p,after=4); run(p,title,True,DARK,13.2)
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8'); bot.set(qn('w:space'),'2'); bot.set(qn('w:color'),TEAL); b.append(bot); pPr.append(b)

def bullet(doc,text,size=8.75):
    p=doc.add_paragraph(style='List Bullet'); para(p,after=.7,line=1.03); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.first_line_indent=Inches(-.14); run(p,text,False,DARK,size)

def role(doc,org,title,period,bullets):
    t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; no_spacing(t); t.columns[0].width=Inches(5.6); t.columns[1].width=Inches(1.5)
    a,b=t.rows[0].cells
    for c in (a,b): cell_margins(c,25,0,10,0)
    p=a.paragraphs[0]; para(p); run(p,org,True,DARK,10.1); run(p,'  |  ',False,MUTED,9.3); run(p,title,True,TEAL,9.4)
    p=b.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; para(p); run(p,period,True,MUTED,8.2)
    for x in bullets: bullet(doc,x)

def add_footer(sec):
    p=sec.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; para(p); run(p,'CASEY BARTON  |  EVIDENCE-BOUND AI SYSTEMS  |  HONOLULU, HAWAII',True,MUTED,6.5)


doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
# Real print-safe margins: visibly present, still compact.
sec.top_margin=Inches(.60); sec.bottom_margin=Inches(.60); sec.left_margin=Inches(.70); sec.right_margin=Inches(.70); sec.header_distance=Inches(.22); sec.footer_distance=Inches(.28)
styles=doc.styles
for sty in ('Normal','List Bullet'):
    styles[sty].font.name=FONT; styles[sty]._element.rPr.rFonts.set(qn('w:eastAsia'),FONT); styles[sty].font.size=Pt(9); styles[sty].font.color.rgb=RGBColor.from_string(DARK)
styles['Normal'].paragraph_format.space_after=Pt(0)
cp=doc.core_properties; cp.title='Casey Barton - Applied AI Systems Architect Resume'; cp.author='Casey Del Carpio Barton'; cp.subject='Evidence-bound applied AI systems resume'; cp.comments='Current claims are separated from ambition and bounded to stated evidence.'

# Header card inside margins
h=doc.add_table(rows=1,cols=2); h.alignment=WD_TABLE_ALIGNMENT.CENTER; h.autofit=False; no_spacing(h); h.columns[0].width=Inches(4.95); h.columns[1].width=Inches(2.15)
a,b=h.rows[0].cells
for c in (a,b): shade(c,NAVY); cell_margins(c,115,125,105,125); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
p=a.paragraphs[0]; para(p,after=2); run(p,'CASEY DEL CARPIO BARTON',True,WHITE,19.2)
p=a.add_paragraph(); para(p,after=2); run(p,'APPLIED AI SYSTEMS ARCHITECT',True,MINT,10.2); run(p,'  |  Agent Infrastructure Engineer',False,'D7E4E7',8.7)
p=a.add_paragraph(); para(p); run(p,'Forward-Deployed AI  |  Builds the operating layer that makes powerful AI dependable enough to use.',False,'C1D1D5',8.1)
p=b.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; para(p,after=2); run(p,'Honolulu, Hawaii',True,WHITE,9.2)
for text,url in [('808-936-5654','tel:+18089365654'),('glacier.equilibrium@gmail.com','mailto:glacier.equilibrium@gmail.com'),('casey-barton-glaciereq.vercel.app','https://casey-barton-glaciereq.vercel.app/'),('github.com/GlacierEQ','https://github.com/GlacierEQ')]:
    p=b.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; para(p,after=0); add_link(p,text,url,'D7E4E7',8.15)

# Signal
s=doc.add_table(rows=1,cols=1); s.alignment=WD_TABLE_ALIGNMENT.CENTER; no_spacing(s); c=s.cell(0,0); shade(c,LIGHT); cell_margins(c,80,105,80,105); borders(c,TEAL,'5')
p=c.paragraphs[0]; para(p,after=2); run(p,'SYSTEMS SIGNAL',True,TEAL,7.4)
p=c.add_paragraph(); para(p,line=1.05); run(p,'Applied AI systems architect who turns ambitious, ambiguous ideas into bounded operating systems with explicit authority, deterministic evidence, controlled failure behavior, and inspectable completion receipts. The method is simple: observe the actual system, separate evidence from inference, isolate the governing failure, build the narrowest useful mechanism, test the stated behavior, expose the limit, and leave a usable artifact.',False,DARK,9.2)

# Metrics
m=doc.add_table(rows=1,cols=4); m.alignment=WD_TABLE_ALIGNMENT.CENTER; m.autofit=False; no_spacing(m)
metrics=[('69/69','Receipt Router tests'),('166 + 19','source + memory tests'),('67','Helix admitted public set'),('62/62','Agent Coordinator tests')]
for i,(num,label) in enumerate(metrics):
    c=m.cell(0,i); shade(c,LIGHT2 if i%2==0 else LIGHT); cell_margins(c,55,35,55,35); borders(c)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; para(p,after=1); run(p,num,True,TEAL,14.2)
    p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; para(p,line=.95); run(p,label,False,MUTED,7.1)

section_title(doc,'Proof at work','Professional Experience')
role(doc,'GlacierEQ','Founder / Applied AI Systems Builder','Jan 2025-Present',[
'Design and implement evidence-bound systems spanning agent orchestration, application intelligence, document and evidence pipelines, memory and continuity, infrastructure governance, and human-machine interfaces.',
'Convert broad architecture into typed contracts, runnable slices, deterministic tests, explicit refusal behavior, machine-readable facts, and reviewable completion receipts.',
'Built a public Portfolio Receipt Router that replaced unsupported autonomous-control semantics with a local fail-closed evidence system; 69/69 tests passed through direct and reusable CI with zero external actions.',
'Govern a multi-repository engineering estate by separating source presence, review, execution, deployment, authority, and verified completion so public claims cannot outrun proof.'
])
role(doc,'Diamond Head Home Inspections','Certified Home Inspector','2020-2024',[
'Inspected residential structural, roofing, electrical, plumbing, HVAC, moisture, and safety conditions; translated field evidence into clear, decision-ready reports under time constraints.',
'Built disciplined habits around uncertainty labeling, defect prioritization, interacting-system analysis, client communication, and avoiding unsupported conclusions.'
])
role(doc,'Hi-Class Home Services / Hi Class Maintenance Oahu LLC','Owner-Operator, Building Systems & Field Services','2017-Present',[
'Scope residential repair and service work, prepare estimates, coordinate execution, and communicate assumptions, constraints, and completion status with clients.',
'Translate incomplete real-world requirements into bounded work packages, explicit assumptions, and usable closeout artifacts.'
])

section_title(doc,'Capability field','Core Capabilities')
cap=doc.add_table(rows=2,cols=2); cap.alignment=WD_TABLE_ALIGNMENT.CENTER; cap.autofit=False; no_spacing(cap); cap.columns[0].width=Inches(3.55); cap.columns[1].width=Inches(3.55)
cap_data=[
('Agent infrastructure & governance','Multi-agent coordination; MCP and tool contracts; capability/authority separation; approval models; state machines; bounded retries; recovery and closure.'),
('Evidence & release engineering','Provenance; claim-to-source mapping; deterministic testing; JSON Schema; CI/CD; static analysis; dependency auditing; integrity hashing; machine-readable receipts.'),
('Application & document intelligence','Role research; evidence-qualified resume generation; document extraction; structured artifacts; human and machine presentation surfaces.'),
('Primary technology','Python; TypeScript; JavaScript; SQL; Bash; Node.js; React; Next.js; FastAPI; REST; JSON-RPC 2.0; MCP; GitHub Actions; Docker; Vercel.')]
for i,(label,val) in enumerate(cap_data):
    c=cap.cell(i//2,i%2); shade(c,LIGHT if i in (0,3) else LIGHT2); cell_margins(c,65,80,65,80); borders(c)
    p=c.paragraphs[0]; para(p,after=2); run(p,label,True,TEAL,8.2)
    p=c.add_paragraph(); para(p,line=1.0); run(p,val,False,DARK,7.75)

# Page 2
p=doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
section_title(doc,'Selected systems','Evidence states, not hype')
rows=[
('Portfolio Receipt Router','TEST VERIFIED','Fail-closed evidence routing and compatibility-safe architecture repair.','69/69 tests; direct and reusable CI; artifact 8910423397; zero external actions.','Independent technical exhibit; no production-operation or company-affiliation claim.'),
('Job Application Helix','PARTIALLY VERIFIED','Evidence-governed hiring and portfolio orchestration with an exact 67-repository admitted public proof boundary.','Fail-closed evidence levels, repository-state separation, promotion gates, machine contracts, and role-calibrated package generation.','The 67 admitted repositories are a public proof set, not accomplishment count or the full owned estate; child repositories retain independent evidence states.'),
('Agent Coordinator','RECORDED 62/62','Deterministic task ownership, dependency order, capacity, stable priority, and shared budgets.','62/62 recorded Python tests.','Hosted matrix and broader promotion gates remain separate.'),
('AKOS / PSYSOC-X','BOUNDED VERIFIED SURFACES','Authority, evidence, closure, and deterministic audience-calibration primitives.','Repository-native verification supports bounded current surfaces.','Presentation depth can change; facts, uncertainty, dignity, and authority may not.'),
]
t=doc.add_table(rows=len(rows),cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; no_spacing(t); t.columns[0].width=Inches(3.25); t.columns[1].width=Inches(3.85)
for i,(name,state,summary,evidence,boundary) in enumerate(rows):
    a,b=t.rows[i].cells
    for c in (a,b): shade(c,LIGHT if i%2==0 else LIGHT2); cell_margins(c,65,70,65,70); borders(c)
    p=a.paragraphs[0]; para(p,after=1); run(p,name,True,DARK,9.7)
    p=a.add_paragraph(); para(p,after=2); run(p,state,True,TEAL if 'BLOCK' not in state else AMBER,7.2)
    p=a.add_paragraph(); para(p,line=1.0); run(p,summary,False,DARK,7.85)
    p=b.paragraphs[0]; para(p,after=1); run(p,'EVIDENCE',True,TEAL,7.0); p=b.add_paragraph(); para(p,after=3,line=1.0); run(p,evidence,False,DARK,7.7)
    p=b.add_paragraph(); para(p,after=1); run(p,'BOUNDARY',True,AMBER,7.0); p=b.add_paragraph(); para(p,line=1.0); run(p,boundary,False,MUTED,7.55)

section_title(doc,'Systems foundation','Cross-domain operating perspective')
f=doc.add_table(rows=3,cols=2); f.alignment=WD_TABLE_ALIGNMENT.CENTER; f.autofit=False; no_spacing(f); f.columns[0].width=Inches(2.0); f.columns[1].width=Inches(5.1)
foundation=[
('Scientific measurement','University of Hawaii sea-urchin morphometric research, 2016-2017: processed physical measurements of gill structures to support visual species identification.'),
('Compressed-gas systems','Scuba Tank Technician, Hawaiian Diving Adventures and UH Dive Office, 2016-2017: tank filling, transport, inspection and repair support, and gas mixing under safety-critical procedures.'),
('Field systems operations','Residential inspection and repair work developed practical judgment around interacting systems, failure consequences, uncertainty, and clear client communication.')]
for i,(label,val) in enumerate(foundation):
    a,b=f.rows[i].cells
    for c in (a,b): shade(c,LIGHT if i%2==0 else LIGHT2); cell_margins(c,50,65,50,65); borders(c)
    p=a.paragraphs[0]; para(p); run(p,label,True,TEAL,7.9)
    p=b.paragraphs[0]; para(p,line=1.0); run(p,val,False,DARK,7.7)

section_title(doc,'Technical profile','Tools and architecture')
p=doc.add_paragraph(); para(p,after=2); run(p,'Primary: ',True,TEAL,8.1); run(p,'Python, TypeScript, JavaScript, SQL, Bash, Node.js, React, Next.js, FastAPI',False,DARK,8.1)
p=doc.add_paragraph(); para(p,after=2); run(p,'Architecture: ',True,TEAL,8.1); run(p,'REST, JSON-RPC 2.0, MCP, JSON Schema, Protocol Buffers, events, state machines, idempotency, authority models, evidence models',False,DARK,8.1)
p=doc.add_paragraph(); para(p,after=2); run(p,'Delivery: ',True,TEAL,8.1); run(p,'Git, GitHub Actions, Docker, Vercel, automated testing, static analysis, dependency auditing, deterministic release checks',False,DARK,8.1)
p=doc.add_paragraph(); para(p); run(p,'Additional evidence-bounded exposure: ',True,TEAL,8.1); run(p,'Go, Rust, C, C++, CUDA, Swift/Metal, Julia, R, Verilog/SystemVerilog, WebAssembly',False,DARK,8.1)

section_title(doc,'Education','Scientific foundation and current cloud training')
e=doc.add_table(rows=2,cols=2); e.alignment=WD_TABLE_ALIGNMENT.CENTER; e.autofit=False; no_spacing(e); e.columns[0].width=Inches(4.4); e.columns[1].width=Inches(2.7)
for i,(school,cred) in enumerate([('University of Hawaii at Manoa','B.S., Marine Biology - 2016'),('AWS Cloud Institute','Cloud Application Developer program - 2025-2026, in progress')]):
    a,b=e.rows[i].cells; cell_margins(a,30,20,30,20); cell_margins(b,30,20,30,20)
    p=a.paragraphs[0]; para(p); run(p,school,True,DARK,8.6); p=b.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; para(p); run(p,cred,False,MUTED,8.0)

# boundary box
box=doc.add_table(rows=1,cols=1); box.alignment=WD_TABLE_ALIGNMENT.CENTER; no_spacing(box); c=box.cell(0,0); shade(c,AMBER_BG); cell_margins(c,65,80,65,80); borders(c,'E5B34B','4')
p=c.paragraphs[0]; para(p,line=1.0); run(p,'EVIDENCE BOUNDARY  ',True,AMBER,7.7); run(p,'Independent GlacierEQ work and bounded technical exhibits. Test counts refer only to their stated repository and scope. No company affiliation, proprietary access, production deployment, customer impact, formal people-management experience, current certification status, hardware validation, or measured business outcome is claimed without direct evidence.',False,DARK,7.65)

for s in doc.sections: add_footer(s)
doc.save(DOCX)
if not DOCX.read_bytes().startswith(b'PK'):
    raise RuntimeError('DOCX signature invalid')
print(DOCX)


# --- HARVESTED FROM {branch} ---
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))

def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def no_cell_spacing(table):
    tblPr = table._tbl.tblPr
    spacing = tblPr.find(qn("w:tblCellSpacing"))
    if spacing is None:
        spacing = OxmlElement("w:tblCellSpacing")
        tblPr.append(spacing)
    spacing.set(qn("w:w"), "0")
    spacing.set(qn("w:type"), "dxa")

def keep_with_next(p, value=True):
    pPr = p._p.get_or_add_pPr()
    node = pPr.find(qn("w:keepNext"))
    if value and node is None:
        pPr.append(OxmlElement("w:keepNext"))
    elif not value and node is not None:
        pPr.remove(node)

def keep_together(p, value=True):
    pPr = p._p.get_or_add_pPr()
    node = pPr.find(qn("w:keepLines"))
    if value and node is None:
        pPr.append(OxmlElement("w:keepLines"))
    elif not value and node is not None:
        pPr.remove(node)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def add_hyperlink(
    paragraph, text, url, color="FFFFFF", underline=False, bold=False, size=8.5
):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    else:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "none")
        rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szcs)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_run(
    p, text, *, bold=False, color=None, size=None, italic=False, all_caps=False
):
    r = p.add_run(text.upper() if all_caps else text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    return r

def set_paragraph(
    p, *, space_before=0, space_after=0, line=1.0, left=0, right=0, keep=False
):
    f = p.paragraph_format
    f.space_before = Pt(space_before)
    f.space_after = Pt(space_after)
    f.line_spacing = line
    f.left_indent = Inches(left) if left else None
    f.right_indent = Inches(right) if right else None
    if keep:
        keep_with_next(p)
        keep_together(p)

def section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, space_before=6, space_after=3, keep=True)
    add_run(p, text, bold=True, color=TEAL_HEX, size=12.2, all_caps=True)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), TEAL_HEX)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

def role_header(doc, org, title, period, location="Honolulu, Hawaii"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.9)
    table.columns[1].width = Inches(1.65)
    no_cell_spacing(table)
    left, right = table.rows[0].cells
    for c in (left, right):
        set_cell_margins(c, top=30, bottom=15, start=0, end=0)
        set_cell_border(c, bottom={"val": "nil"})
    p = left.paragraphs[0]
    set_paragraph(p, space_after=0)
    add_run(p, org, bold=True, color=DARK_HEX, size=10.7)
    add_run(p, " | ", color=MUTED_HEX, size=10.5)
    add_run(p, title, bold=True, color=TEAL_HEX, size=10.3)
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(p2, space_after=0)
    add_run(p2, period, bold=True, color=MUTED_HEX, size=8.7)
    lp = doc.add_paragraph()
    set_paragraph(lp, space_after=1)
    add_run(lp, location, italic=True, color=MUTED_HEX, size=8.6)
    return table

def add_label_value(
    cell, label, value, label_color=TEAL_HEX, value_color=DARK_HEX, value_size=8.35
):
    p = cell.add_paragraph()
    set_paragraph(p, space_after=1)
    add_run(p, label, bold=True, color=label_color, size=7.25, all_caps=True)
    p2 = cell.add_paragraph()
    set_paragraph(p2, space_after=2, line=1.02)
    add_run(p2, value, color=value_color, size=value_size)

def set_doc_core(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.42)
    sec.bottom_margin = Inches(0.45)
    sec.left_margin = Inches(0.48)
    sec.right_margin = Inches(0.48)
    sec.header_distance = Inches(0.15)
    sec.footer_distance = Inches(0.25)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(9.2)
    styles["Normal"].font.color.rgb = RGBColor.from_string(DARK_HEX)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    # Metadata
    cp = doc.core_properties
    cp.title = "Casey Barton - Applied AI Systems Architect Resume V17"
    cp.subject = "PSYSOC-X calibrated executive resume with machine-verifiable evidence boundaries"
    cp.author = "Casey Del Carpio Barton"
    cp.keywords = "Applied AI Systems Architect, Agent Infrastructure, MCP, multi-agent systems, evidence governance, deterministic validation, Python, TypeScript, Forward-Deployed AI"
    cp.comments = "Generated from a canonical evidence-bound resume model; claims are scoped to stated receipts."

def make_docx() -> None:
    doc = Document()
    set_doc_core(doc)

    # HEADER BAND
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Inches(5.55)
    header.columns[1].width = Inches(2.0)
    no_cell_spacing(header)
    left, right = header.rows[0].cells
    for c in (left, right):
        shade(c, NAVY_HEX)
        set_cell_margins(c, top=125, bottom=115, start=115, end=115)
        set_cell_border(
            c,
            top={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
            bottom={"val": "nil"},
        )
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    set_paragraph(p, space_after=2)
    add_run(p, "CASEY DEL CARPIO BARTON", bold=True, color=WHITE_HEX, size=20.5)
    p = left.add_paragraph()
    set_paragraph(p, space_after=2)
    add_run(p, "APPLIED AI SYSTEMS ARCHITECT", bold=True, color=MINT_HEX, size=10.5)
    add_run(p, "  |  Agent Infrastructure Engineer", color="D7E4E7", size=9.0)
    p = left.add_paragraph()
    set_paragraph(p, space_after=0)
    add_run(
        p,
        "Forward-Deployed AI  |  Evidence-bound execution for ambitious systems and high-consequence decisions.",
        color="C1D1D5",
        size=8.15,
    )
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(p, space_after=1)
    add_run(p, "Honolulu, Hawaii", bold=True, color=WHITE_HEX, size=9.4)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(p, space_after=0)
    add_hyperlink(p, "808-936-5654", "tel:+18089365654", color="D7E4E7", size=8.4)
    add_run(p, "\n", size=1)
    add_hyperlink(
        p,
        "glacier.equilibrium@gmail.com",
        "mailto:glacier.equilibrium@gmail.com",
        color="D7E4E7",
        size=8.4,
    )
    add_run(p, "\n", size=1)
    add_hyperlink(
        p,
        "casey-barton-glaciereq.vercel.app",
        "https://casey-barton-glaciereq.vercel.app/",
        color="D7E4E7",
        size=8.4,
    )
    add_run(p, "\n", size=1)
    add_hyperlink(
        p,
        "github.com/GlacierEQ",
        "https://github.com/GlacierEQ",
        color="D7E4E7",
        size=8.4,
    )

    # EXECUTIVE SIGNAL
    signal = doc.add_table(rows=1, cols=1)
    signal.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_cell_spacing(signal)
    c = signal.cell(0, 0)
    shade(c, LIGHT_HEX)
    set_cell_margins(c, top=75, bottom=75, start=90, end=90)
    set_cell_border(
        c,
        top={"val": "single", "sz": "8", "color": TEAL_HEX},
        bottom={"val": "single", "sz": "8", "color": TEAL_HEX},
        left={"val": "single", "sz": "2", "color": MID_HEX},
        right={"val": "single", "sz": "2", "color": MID_HEX},
    )
    p = c.paragraphs[0]
    set_paragraph(p, space_after=2)
    add_run(p, "EXECUTIVE SIGNAL", bold=True, color=TEAL_HEX, size=8.2, all_caps=True)
    p = c.add_paragraph()
    set_paragraph(p, space_after=0, line=1.06)
    add_run(
        p,
        "Builds the operating layer between model capability and dependable outcomes: explicit authority, bounded tool use, deterministic evidence, controlled failure, continuity, and completion receipts. Combines software architecture with scientific measurement, compressed-gas safety, building-system inspection, and owner-operated field execution.",
        color=DARK_HEX,
        size=9.7,
    )

    # METRICS
    metrics = doc.add_table(rows=1, cols=4)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics.autofit = False
    no_cell_spacing(metrics)
    for idx, (num, label) in enumerate(
        [
            ("69/69", "Receipt Router tests"),
            ("166 + 19", "Bounded source + memory tests"),
            ("148/148", "Job Application Helix tests"),
            ("0", "External actions in flagship verification"),
        ]
    ):
        cell = metrics.cell(0, idx)
        shade(cell, LIGHT2_HEX if idx % 2 == 0 else LIGHT_HEX)
        set_cell_margins(cell, top=65, bottom=65, start=40, end=40)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "2", "color": MID_HEX},
            bottom={"val": "single", "sz": "2", "color": MID_HEX},
            left={"val": "single", "sz": "2", "color": MID_HEX},
            right={"val": "single", "sz": "2", "color": MID_HEX},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, space_after=1)
        add_run(p, num, bold=True, color=TEAL_HEX, size=15.8)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, space_after=0, line=0.95)
        add_run(p, label, color=MUTED_HEX, size=7.5)

    section_heading(doc, "Professional Experience")
    role_header(
        doc, "GlacierEQ", "Founder / Applied AI Systems Builder", "Jan 2025-Present"
    )
    for t in [
        "Design and implement evidence-bound systems spanning agent orchestration, application intelligence, document and evidence pipelines, memory and continuity, infrastructure governance, and human-machine interfaces.",
        "Convert ambiguous architecture into typed contracts, runnable slices, deterministic tests, explicit refusal behavior, machine-readable facts, and reviewable completion receipts.",
        "Built a public Portfolio Receipt Router that replaced unsupported autonomous-control semantics with a local fail-closed evidence system; 69 of 69 tests passed through direct and reusable CI with zero external actions.",
        "Govern a multi-repository portfolio by separating source presence, review, execution, deployment, authority, and verified completion so recruiter, technical, and machine views cannot drift.",
    ]:
        bullet(doc, t, size=9.0)

    role_header(
        doc, "Diamond Head Home Inspections", "Certified Home Inspector", "2020-2024"
    )
    for t in [
        "Inspected residential structural, roofing, electrical, plumbing, HVAC, moisture, and safety conditions; translated field evidence into clear, decision-ready reports under time constraints.",
        "Developed disciplined habits around uncertainty labeling, defect prioritization, interacting-system analysis, client communication, and avoiding unsupported conclusions.",
    ]:
        bullet(doc, t, size=9.0)

    role_header(
        doc,
        "Hi-Class Home Services / Hi Class Maintenance Oahu LLC",
        "Owner-Operator, Building Systems & Field Services",
        "2017-Present",
    )
    for t in [
        "Scope residential repair and service work, prepare estimates, coordinate execution, and communicate assumptions, constraints, and completion status with clients.",
        "Work across practical building-system problems including minor carpentry, plumbing, electrical fixture replacement, drywall, painting, flooring, and maintenance planning.",
    ]:
        bullet(doc, t, size=9.0)

    section_heading(doc, "Core Capability Matrix")
    cap = doc.add_table(rows=2, cols=2)
    cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    cap.autofit = False
    no_cell_spacing(cap)
    cap_data = [
        (
            "Agent infrastructure & governance",
            "Multi-agent coordination; MCP and tool contracts; capability/authority separation; approval models; state machines; bounded retries; recovery and closure.",
        ),
        (
            "Evidence & release engineering",
            "Provenance; claim-to-source mapping; deterministic testing; JSON Schema; CI/CD; static analysis; integrity hashing; machine-readable receipts.",
        ),
        (
            "Application & document intelligence",
            "Role research; evidence-qualified resume generation; document extraction; structured artifacts; human and machine presentation layers.",
        ),
        (
            "Primary technology",
            "Python; TypeScript; JavaScript; SQL; Bash; Node.js; React; Next.js; FastAPI; REST; JSON-RPC 2.0; MCP; GitHub Actions; Docker; Vercel.",
        ),
    ]
    for i, (label, value) in enumerate(cap_data):
        cell = cap.cell(i // 2, i % 2)
        shade(cell, LIGHT2_HEX if i % 3 else LIGHT_HEX)
        set_cell_margins(cell, top=65, bottom=65, start=80, end=80)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "2", "color": MID_HEX},
            bottom={"val": "single", "sz": "2", "color": MID_HEX},
            left={"val": "single", "sz": "2", "color": MID_HEX},
            right={"val": "single", "sz": "2", "color": MID_HEX},
        )
        p = cell.paragraphs[0]
        set_paragraph(p, space_after=2)
        add_run(p, label, bold=True, color=TEAL_HEX, size=8.3)
        p = cell.add_paragraph()
        set_paragraph(p, space_after=0, line=1.0)
        add_run(p, value, color=DARK_HEX, size=7.8)

    # page break
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    section_heading(doc, "Selected Systems and Evidence States")
    sys_table = doc.add_table(rows=5, cols=2)
    sys_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sys_table.autofit = False
    no_cell_spacing(sys_table)
    sys_table.columns[0].width = Inches(3.9)
    sys_table.columns[1].width = Inches(3.65)
    systems = [
        (
            "Portfolio Receipt Router",
            "TEST VERIFIED",
            "Compatibility-safe architecture repair and evidence routing.",
            "69/69 tests; direct and reusable CI; artifact 8910423397; zero external actions.",
            "Independent technical exhibit; no datacenter control, company affiliation, or production-operation claim.",
        ),
        (
            "Job Application Helix",
            "RECORDED 148/148",
            "Evidence-governed hiring and portfolio orchestration.",
            "148/148 recorded repository tests.",
            "Release-specific current-SHA and deployment gates remain separate.",
        ),
        (
            "Agent Coordinator",
            "RECORDED 62/62",
            "Deterministic task ownership, dependencies, capacity, priority, and shared budgets.",
            "62/62 recorded Python tests.",
            "Hosted matrix and build-wheel promotion remain open.",
        ),
        (
            "Microcode Governance",
            "REVIEWED - EXECUTION BLOCKED",
            "Firmware manifests, drift, SBOMs, provenance, compatibility policy, and approval-gated rollout planning.",
            "132-check generated contract; static review found no critical defect.",
            "Private CI did not execute; not test verified and performs no firmware mutation.",
        ),
        (
            "PSYSOC-X",
            "BOUNDED TEST SCOPE",
            "Audience calibration across recruiter, master, machine, and relationship views.",
            "Deterministic profile and invariant cases.",
            "No diagnosis, covert persuasion, or manipulation capability is claimed.",
        ),
    ]
    for i, (name, state, summary, evidence, boundary) in enumerate(systems):
        left, right = sys_table.rows[i].cells
        for cell in (left, right):
            shade(cell, LIGHT_HEX if i % 2 == 0 else LIGHT2_HEX)
            set_cell_margins(cell, top=65, bottom=65, start=65, end=65)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "2", "color": MID_HEX},
                bottom={"val": "single", "sz": "2", "color": MID_HEX},
                left={"val": "single", "sz": "2", "color": MID_HEX},
                right={"val": "single", "sz": "2", "color": MID_HEX},
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = left.paragraphs[0]
        set_paragraph(p, space_after=1)
        add_run(p, name, bold=True, color=DARK_HEX, size=10.2)
        p = left.add_paragraph()
        set_paragraph(p, space_after=2)
        add_run(
            p,
            state,
            bold=True,
            color=AMBER_HEX if "BLOCKED" in state else TEAL_HEX,
            size=7.3,
        )
        p = left.add_paragraph()
        set_paragraph(p, space_after=0, line=1.0)
        add_run(p, summary, color=DARK_HEX, size=8.25)
        add_label_value(
            right, "Evidence", evidence, label_color=TEAL_HEX, value_size=8.05
        )
        add_label_value(
            right,
            "Boundary",
            boundary,
            label_color=AMBER_HEX,
            value_color=MUTED_HEX,
            value_size=7.9,
        )

    section_heading(doc, "Cross-Domain Systems Foundation")
    found = doc.add_table(rows=4, cols=2)
    found.alignment = WD_TABLE_ALIGNMENT.CENTER
    found.autofit = False
    no_cell_spacing(found)
    found.columns[0].width = Inches(2.1)
    found.columns[1].width = Inches(5.45)
    found_data = [
        (
            "Scientific measurement",
            "University of Hawaii sea-urchin morphometric research, 2016-2017: processed physical measurements of gill structures to support visual species identification.",
        ),
        (
            "Compressed-gas systems",
            "Scuba Tank Technician, Hawaiian Diving Adventures and UH Dive Office, 2016-2017: tank filling, transport, inspection and repair support, and gas mixing under safety-critical procedures.",
        ),
        (
            "Field systems operations",
            "Residential inspection and repair work developed practical judgment around interacting systems, failure consequences, uncertainty, and clear client communication.",
        ),
        (
            "Earlier leadership",
            "Civil Air Patrol Cadet Commander and Search-and-Rescue Ground Team Leader, 2007-2011; included as historical leadership context, not as formal professional management experience.",
        ),
    ]
    for i, (label, value) in enumerate(found_data):
        c1, c2 = found.rows[i].cells
        for c in (c1, c2):
            shade(c, LIGHT_HEX if i % 2 == 0 else LIGHT2_HEX)
            set_cell_margins(c, top=55, bottom=55, start=65, end=65)
            set_cell_border(
                c,
                top={"val": "single", "sz": "2", "color": MID_HEX},
                bottom={"val": "single", "sz": "2", "color": MID_HEX},
                left={"val": "single", "sz": "2", "color": MID_HEX},
                right={"val": "single", "sz": "2", "color": MID_HEX},
            )
        p = c1.paragraphs[0]
        set_paragraph(p, space_after=0)
        add_run(p, label, bold=True, color=TEAL_HEX, size=8.3)
        p = c2.paragraphs[0]
        set_paragraph(p, space_after=0, line=1.0)
        add_run(p, value, color=DARK_HEX, size=8.15)

    section_heading(doc, "Education and Earlier Technical Credentials")
    edu = doc.add_table(rows=2, cols=2)
    edu.alignment = WD_TABLE_ALIGNMENT.CENTER
    edu.autofit = False
    no_cell_spacing(edu)
    edu.columns[0].width = Inches(4.8)
    edu.columns[1].width = Inches(2.75)
    for i, (school, credential) in enumerate(
        [
            ("University of Hawaii at Manoa", "B.S., Marine Biology - 2016"),
            (
                "AWS Cloud Institute",
                "Cloud Application Developer program - 2025-2026, in progress",
            ),
        ]
    ):
        c1, c2 = edu.rows[i].cells
        for c in (c1, c2):
            set_cell_margins(c, top=35, bottom=35, start=20, end=20)
            set_cell_border(c, bottom={"val": "nil"})
        p = c1.paragraphs[0]
        set_paragraph(p, space_after=0)
        add_run(p, school, bold=True, color=DARK_HEX, size=9.0)
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph(p, space_after=0)
        add_run(p, credential, color=MUTED_HEX, size=8.5)

    p = doc.add_paragraph()
    set_paragraph(p, space_before=4, space_after=4, line=1.0)
    add_run(
        p,
        "Earlier technical certifications recorded in prior resumes: ",
        bold=True,
        color=TEAL_HEX,
        size=7.85,
    )
    add_run(
        p,
        "PSI Visual Cylinder Inspector, Eddy Current Technician, Valve Repair Technician, Oxygen Cleaning Cylinder Technician; NAUI Rescue and Master Diver; PADI Enriched Air Diver. Current status is not represented as active without updated documentation.",
        color=MUTED_HEX,
        size=7.7,
    )

    # evidence boundary box
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_cell_spacing(box)
    c = box.cell(0, 0)
    shade(c, AMBER_BG_HEX)
    set_cell_margins(c, top=65, bottom=65, start=75, end=75)
    set_cell_border(
        c,
        top={"val": "single", "sz": "4", "color": "E5B34B"},
        bottom={"val": "single", "sz": "4", "color": "E5B34B"},
        left={"val": "single", "sz": "2", "color": "F2D99A"},
        right={"val": "single", "sz": "2", "color": "F2D99A"},
    )
    p = c.paragraphs[0]
    set_paragraph(p, space_after=0, line=1.0)
    add_run(p, "EVIDENCE BOUNDARY  ", bold=True, color=AMBER_HEX, size=8.0)
    add_run(
        p,
        "Independent GlacierEQ work and bounded technical exhibits. Test counts refer only to their stated repository and scope. No company affiliation, proprietary access, production use, customer impact, formal people-management experience, current certification status, or hardware validation is claimed without direct evidence.",
        color=DARK_HEX,
        size=7.9,
    )

    # Save and scrub metadata later
    for sec in doc.sections:
        add_footer(sec)
    doc.save(DOCX_PATH)

def main() -> None:
    make_docx()
    if not DOCX_PATH.read_bytes().startswith(b"PK"):
        raise RuntimeError("DOCX signature invalid")
    print(
        json.dumps(
            {
                "docx": {
                    "bytes": DOCX_PATH.stat().st_size,
                    "sha256": hashlib.sha256(DOCX_PATH.read_bytes()).hexdigest(),
                },
                "pdf": {
                    "mode": "committed visual export; identity bound by resume-artifacts.json"
                },
            },
            indent=2,
            sort_keys=True,
        )
    )



# --- HARVESTED FROM {branch} ---
def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()

def pdf_styles():
    s = getSampleStyleSheet()
    return {
        "name": ParagraphStyle("name", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=28, leading=29, textColor=WHITE, spaceAfter=5),
        "title": ParagraphStyle("title", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=10.2, leading=13, textColor=MINT),
        "contact": ParagraphStyle("contact", parent=s["Normal"], fontName="Helvetica", fontSize=7.7, leading=10.2, textColor=colors.HexColor("#D4E6E2")),
        "section": ParagraphStyle("section", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=12.4, leading=14, textColor=INK, spaceBefore=5, spaceAfter=6),
        "tag": ParagraphStyle("tag", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=6.3, leading=8, textColor=TEAL, spaceAfter=3),
        "body": ParagraphStyle("body", parent=s["Normal"], fontName="Helvetica", fontSize=8.15, leading=11.1, textColor=INK),
        "small": ParagraphStyle("small", parent=s["Normal"], fontName="Helvetica", fontSize=7.1, leading=9.5, textColor=MUTED),
        "card_title": ParagraphStyle("card_title", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=9, leading=10.5, textColor=INK, spaceAfter=4),
        "card_state": ParagraphStyle("card_state", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=5.8, leading=7, textColor=TEAL, spaceAfter=5),
        "metric": ParagraphStyle("metric", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=16.5, leading=17, textColor=MINT, alignment=TA_LEFT),
        "metric_label": ParagraphStyle("metric_label", parent=s["Normal"], fontName="Helvetica", fontSize=6.5, leading=8, textColor=colors.HexColor("#D8E8E5")),
        "role": ParagraphStyle("role", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=8.6, leading=10, textColor=INK),
        "date": ParagraphStyle("date", parent=s["Normal"], fontName="Helvetica-Bold", fontSize=7.2, leading=9, textColor=TEAL, alignment=TA_LEFT),
        "bullet": ParagraphStyle("bullet", parent=s["Normal"], fontName="Helvetica", fontSize=7.6, leading=10.3, leftIndent=10, firstLineIndent=-7, bulletIndent=0, textColor=INK),
        "boundary": ParagraphStyle("boundary", parent=s["Normal"], fontName="Helvetica", fontSize=6.7, leading=9, textColor=colors.HexColor("#5F5032")),
    }

def make_pdf() -> None:
    styles = pdf_styles()
    doc = BaseDocTemplate(str(PDF_PATH), pagesize=letter, leftMargin=0.42 * inch, rightMargin=0.42 * inch, topMargin=0.37 * inch, bottomMargin=0.34 * inch, title="Casey Barton - Applied AI Systems Architect Resume", author="Casey Del Carpio Barton")
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    doc.addPageTemplates([PageTemplate(id="resume", frames=[frame])])
    story = []

    header_left = [Paragraph("CASEY DEL CARPIO BARTON", styles["name"]), Paragraph("APPLIED AI SYSTEMS ARCHITECT | AGENT INFRASTRUCTURE ENGINEER", styles["title"]), Spacer(1, 6), Paragraph("Builds the operating layer that makes powerful AI dependable enough to use.", styles["contact"])]
    header_right = [Paragraph("Honolulu, Hawaii", styles["contact"]), Paragraph("glacier.equilibrium@gmail.com", styles["contact"]), Paragraph("casey-barton-glaciereq.vercel.app", styles["contact"]), Paragraph("github.com/GlacierEQ", styles["contact"])]
    header = Table([[header_left, header_right]], colWidths=[4.85 * inch, 2.05 * inch], hAlign="LEFT")
    header.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), DARK), ("BOX", (0, 0), (-1, -1), 0.6, DARK_2), ("VALIGN", (0, 0), (-1, -1), "BOTTOM"), ("LEFTPADDING", (0, 0), (0, 0), 18), ("RIGHTPADDING", (1, 0), (1, 0), 16), ("TOPPADDING", (0, 0), (-1, -1), 17), ("BOTTOMPADDING", (0, 0), (-1, -1), 16)]))
    story.extend([header, Spacer(1, 10)])

    metric_cells = []
    for value, label in PROOF:
        metric_cells.append([Paragraph(value, styles["metric"]), Paragraph(label, styles["metric_label"])])
    metrics = Table([metric_cells], colWidths=[doc.width / 4] * 4)
    metrics.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), DARK_2), ("BOX", (0, 0), (-1, -1), 0.5, DARK_2), ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#2D6662")), ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
    story.extend([metrics, Spacer(1, 10), Paragraph("SYSTEMS POSITIONING", styles["tag"]), Paragraph("Inspection discipline for AI systems.", styles["section"]), Paragraph(SUMMARY, styles["body"]), Spacer(1, 8)])

    project_cells = []
    for p in PROJECTS:
        state_color = AMBER if p["state"] == "EXECUTION BLOCKED" else TEAL
        block = [Paragraph(p["state"], ParagraphStyle("state", parent=styles["card_state"], textColor=state_color)), Paragraph(p["name"], styles["card_title"]), Paragraph(p["text"], styles["small"]), Spacer(1, 4), Paragraph(p["proof"], styles["small"])]
        project_cells.append(block)
    project_table = Table([[project_cells[0], project_cells[1]], [project_cells[2], project_cells[3]]], colWidths=[doc.width / 2 - 4, doc.width / 2 - 4], hAlign="LEFT")
    project_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.white), ("BOX", (0, 0), (-1, -1), 0.5, LINE), ("INNERGRID", (0, 0), (-1, -1), 0.5, LINE), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 11), ("RIGHTPADDING", (0, 0), (-1, -1), 11), ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9), ("BACKGROUND", (1, 1), (1, 1), AMBER_BG)]))
    story.extend([Paragraph("SELECTED EXECUTION", styles["tag"]), Paragraph("Proof that reveals engineering judgment.", styles["section"]), project_table, Spacer(1, 10), Paragraph("EXPERIENCE", styles["tag"])])

    first = EXPERIENCE[0]
    exp_head = Table([[Paragraph(f"{first[0]} - {first[1]}", styles["role"]), Paragraph(first[2], styles["date"])]], colWidths=[5.35 * inch, 1.55 * inch])
    exp_head.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("ALIGN", (1, 0), (1, 0), "RIGHT"), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    story.append(exp_head)
    for b in first[3]:
        story.append(Paragraph(f"- {b}", styles["bullet"]))

    story.append(PageBreak())
    story.extend([Paragraph("EXPERIENCE CONTINUED", styles["tag"])])
    for org, role, period, bullets in EXPERIENCE[1:]:
        head = Table([[Paragraph(f"{org} - {role}", styles["role"]), Paragraph(period, styles["date"])]], colWidths=[5.35 * inch, 1.55 * inch])
        head.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("ALIGN", (1, 0), (1, 0), "RIGHT"), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
        story.append(head)
        for b in bullets:
            story.append(Paragraph(f"- {b}", styles["bullet"]))

    story.extend([Spacer(1, 8), Paragraph("CROSS-DOMAIN FOUNDATION", styles["tag"]), Paragraph("A rare operating perspective.", styles["section"])])
    domain_cells = []
    for tag, title, text in DOMAINS:
        domain_cells.append([Paragraph(tag, styles["card_state"]), Paragraph(title, styles["card_title"]), Paragraph(text, styles["small"])])
    domains = Table([domain_cells], colWidths=[doc.width / 3] * 3)
    domains.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), PALE), ("BOX", (0, 0), (-1, -1), 0.5, LINE), ("INNERGRID", (0, 0), (-1, -1), 0.5, LINE), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 9), ("RIGHTPADDING", (0, 0), (-1, -1), 9), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
    story.extend([domains, Spacer(1, 8), Paragraph("CAPABILITIES", styles["tag"])])
    cap_cells = []
    for title, text in CAPABILITIES:
        cap_cells.append([Paragraph(title, styles["card_title"]), Paragraph(text, styles["small"])])
    caps = Table([cap_cells], colWidths=[doc.width / 3] * 3)
    caps.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.white), ("BOX", (0, 0), (-1, -1), 0.5, LINE), ("INNERGRID", (0, 0), (-1, -1), 0.5, LINE), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 9), ("RIGHTPADDING", (0, 0), (-1, -1), 9), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
    story.extend([caps, Spacer(1, 8)])

    edu = Table([[Paragraph("University of Hawaii at Manoa", styles["role"]), Paragraph("B.S., Marine Biology - 2016", styles["small"]), Paragraph("AWS Cloud Institute", styles["role"]), Paragraph("Cloud Application Developer program - 2025-2026, in progress", styles["small"])]], colWidths=[1.55 * inch, 1.75 * inch, 1.35 * inch, 2.25 * inch])
    edu.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), PALE), ("BOX", (0, 0), (-1, -1), 0.5, LINE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
    story.extend([Paragraph("EDUCATION", styles["tag"]), edu, Spacer(1, 7), Paragraph("Earlier technical certifications recorded in prior resumes: PSI Visual Cylinder Inspector, Eddy Current Technician, Valve Repair Technician, Oxygen Cleaning Cylinder Technician; NAUI Rescue and Master Diver; PADI Enriched Air Diver. Current status should be confirmed before role-specific use.", styles["small"]), Spacer(1, 7)])
    boundary = Table([[Paragraph(f"<b>Evidence boundary.</b> {BOUNDARY}", styles["boundary"])]], colWidths=[doc.width])
    boundary.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), AMBER_BG), ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E7C77E")), ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
    story.append(boundary)
    doc.build(story)

def add_heading(doc: Document, tag: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, tag.upper(), 7, True, "0D766C")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(5)
    add_run(p2, title, 15, True, "12272D")

def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.5)
    add_run(p, text, 8.2, False, "263C41")

